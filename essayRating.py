# Built-in libraries
import traceback
import pypandoc
# Third-party libraries
import openai
from flask import request, jsonify, send_file, make_response
from sqlalchemy.orm import sessionmaker
from flask import Blueprint, abort
from contextlib import contextmanager
from flask_jwt_extended import jwt_required, get_jwt_identity
from docx import Document
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
import re
import ast
import openai
# Local modules
from my_extensions import db
from functions import load_sensitive_words, contains_sensitive_words, update_generate_count, generate_pdf,check_generate_count


essayRating_blueprint = Blueprint("essayRating", __name__)

openai.api_key = "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"
openai.Model.list()

@contextmanager
def create_session():
    """Provide a transactional scope around a series of operations."""
    Session = sessionmaker(bind=db.engine)
    session = Session()
    try:
        yield session
        session.commit()
    except:
        session.rollback()
        raise
    finally:
        session.close()
        
        
def count_words(text):
    # Use the regex \w+ to match words
    words = re.findall(r'\w+', text)
    return len(words)

def diff_strings(str1, str2):
    diff_index = [(m.start(), m.end()) for m in re.finditer(r'(?<=[^\w])' + re.escape(str2) + r'(?=[^\w])', str1)]
    if diff_index:
        return diff_index[0]
    else:
        return None, None
#/////////////////////////#///////////////////////////#


#------------------ generate_ielts_report --------------------------#

@essayRating_blueprint.route("/api/generate_ielts_report", methods=["POST"])
@jwt_required()
@check_generate_count(40)
def generate_ielts_report():
    
    #define IeltsWritingScoreCritiria & Sampleanswers
    IeltsWritingScoreCritiria = f'''
        Task Response

        7: The principal parts of the task are correctly addressed. A clear and developed position is presented. Main ideas are expanded and supported, though there might be a tendency to overgeneralize or lack focus and precision in supporting ideas/material.
        6: The main aspects of the prompt are addressed, though some aspects are more detailed than others. A pertinent position is presented but conclusions might be unclear, unjustified, or repetitive. Main ideas are relevant, but some might lack development or clarity.
        5: The main aspects of the prompt are only partially addressed and might be inappropriate in places. A position is presented but the development might be unclear. Main ideas might be limited and insufficiently developed.
        4: The prompt is minimally tackled or the response might be tangential due to some misunderstanding. A position is present but might require careful reading to be understood.
        3: The prompt is inadequately addressed or misunderstood. No relevant position can be identified. Ideas might be sparse and insufficiently developed.
        2: The content barely relates to the prompt. No position can be identified. Ideas are barely discernible and lack development.
        1: The content is unrelated to the prompt. No position can be identified. Any copied rubric must be discounted.

        Coherence & Cohesion

        7: Information and ideas are logically organized, and there is a clear progression throughout the response. A range of cohesive devices is used flexibly, but with some inaccuracies or some over/underuse. Paragraphing is generally used effectively.
        6: Information and ideas are generally arranged coherently with clear overall progression. Cohesive devices are used to good effect, but there might be misuse or overuse.
        5: Organization is evident but not completely logical. There might be a lack of overall progression. The relationship between ideas can be followed but sentences are not fluently linked.
        4: Information and ideas are not arranged coherently and there's no clear progression. Relationships between ideas can be unclear. Paragraphing may be absent.
        3: There's no apparent logical organization. Ideas are discernible but difficult to relate to each other. Any attempts at paragraphing are unhelpful.
        2: There's little relevant message, or the entire response may be off-topic. There's little evidence of control of organizational features.
        1: The writing fails to communicate any message and appears to be by a virtual non-writer. Responses of 20 words or fewer are rated at Band 1.

        Lexical Resource

        7: The resource is sufficient to allow some flexibility and precision. There is some ability to use less common and/or idiomatic items. An awareness of style and collocation is evident, though inaccuracies occur.
        6: The resource is generally adequate for the task. The meaning is generally clear despite a restricted range or a lack of precision in word choice.
        5: The resource is limited but minimally adequate for the task. Simple vocabulary might be used accurately but lacks variation in expression.
        4: The resource is limited and inadequate for or unrelated to the task. Vocabulary is basic and used repetitively.
        3: The resource is inadequate. Control of word choice and spelling is very limited, and errors predominate.
        2: The resource is extremely limited with few recognizable strings, apart from memorized phrases.
        1: No resource is apparent, except for a few isolated words. Responses of 20 words or fewer are rated at Band 1.

        Grammatical Range & Accuracy

        7: A variety of complex structures is used with some flexibility and accuracy. Grammar and punctuation are generally well-controlled, and error-free sentences are frequent. A few grammar errors may persist, but these do not impede communication.
        6: A mix of simple and complex sentence forms is used but flexibility is limited. Errors in grammar and punctuation occur but rarely impede communication.
        5: The range of structures is limited and repetitive. Grammatical errors might be frequent and cause some difficulty for the reader. Punctuation might be faulty.
        4: A very limited range of structures is used. Grammatical errors are frequent and may impede meaning.
        3: Sentence forms are attempted, but errors in grammar and punctuation predominate. Length may be insufficient to provide evidence of control of sentence forms.
        2: There's little or no evidence of sentence forms. There is no apparent control of word formation and spelling.

            '''
    Sampleanswers = f'''
        IELTS essay sample question (1)
        Many students are studying abroad as part of their degree programme. What are the advantages and disadvantages
        
        IELTS BAND 6 sample essay:
        Nowadays, It is getting popular that study for a semester in foreign countries as an international student to build a better profile in their carrier whereas many students are going abroad for studies in between degrees. I am going to discuss the positives and negatives of studying in cross countries. There are quite a lot of advantages in this scenario for the individual.

        Firstly, pursuing a semester in an international university gives better opportunities in building one career in society. Coming the positive, they can have a better knowledge about the institutes and their course structures of other countries whereas very helpful for students who want to pursue their higher studies in international organizations. Secondly, they come to know the culture and customs of the country. For example, most universities are accepting overseas students to study on their campuses for a semester.

        In contrast, going abroad in between undergraduate program study tends to spend a big amount as the universities will not provide any scholarships or incentives to students who came from other countries. Moreover, there is a problem in coping  with students as they will be more advance in different types of skills like designing, analyzing problems and so on. For instance, some of the students who went from India to study abroad faced fewer problems in competing with them.

        In conclusion, people who are going to abroad for studying as a part of their degree program have some positives and negatives. There are quite a lot of advantages like gaining better knowledge of universities and their curriculums and also their culture. On the other hand, the idea of studying in this way  totends spend a huge amount. These all depend on people who want to study like that in other countries.
                '''
                
    try:
        data = request.get_json()
        essay_question = data.get("essay_question")
        essay = data.get("essay")
        


        # 加载敏感词
        sensitive_words = load_sensitive_words('chinese')

        # 检查是否包含敏感词
        contains_sensitive = contains_sensitive_words(essay_question, sensitive_words)
        if contains_sensitive:
            return jsonify({"error": "您的输入包含敏感词，请重新输入。"}), 400
        
        contains_sensitive = contains_sensitive_words(essay, sensitive_words)
        if contains_sensitive:
            return jsonify({"error": "您的输入包含敏感词，请重新输入。"}), 400
        
        def get_correctionsandimprovements(essay):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    你是一个资深的雅思考官及python 专家，你的任务是为一篇雅思大作文提供专业的修改建议和改进建议。请首先阅读这篇文章，根据雅思的标准和特点生成三个词典：

                    1. `corrections`：这是一个字典，键是文章中需要修正的部分，修改词汇用错、语法、标点符号用错的地方，值是你用中文写的建议的修正和修正理由。
                    2. `improvements`：这是一个字典，键是文章中可以改进的部分主要集中于Coherence & Cohesion, diction 和 sentence variety，值是你用中文写的建议的改进理由。
                    3. `diction`：这是一个字典，键是文章中可以paraphrase 的词汇，主要集中于用词重复或者不准确，值是你用建议的单词。
                    三个字典都要列出5项。
                    
                    
                    你要写纯粹的字典格式，并且不包括任何其他说明或者引用符号。格式和下面的例子保持100%一致：
                    比如原文有一句“The speed of the car is high”，那么corrections和improvements修改格式如下：
                    corrections = {
                        "is fast" ':' "is high”在此处使用不当，建议改为The speed of the car is high”。因为..."
                    }
                    improvements = {
                        "The speed of the car is fast." ':' "建议改为The speed of the car is high,因为..."
                    }
                    diction = {
                        "increse." ':' "建议改为rise"
                    }
                    
                    """
                    现在请你从阅读用户作文开始。我不需要任何警告和道歉。

                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    {essay}
                    
                    '''
            }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=messages,
                temperature=0,
                max_tokens=1000,
                #stream=True
            )
            string = response.choices[0].message.content.strip()
                # 删除 Python 关键字
            response_string = string.replace("corrections = ", "").replace("improvements = ", "").replace("diction = ", "")
            # 将字符串分割成两部分，每部分包含一个字典
            dict_strings = response_string.split("\n\n")
            # 对每一部分使用 ast.literal_eval() 函数，得到两个字典
            corrections = ast.literal_eval(dict_strings[0])
            improvements = ast.literal_eval(dict_strings[1])
            diction = ast.literal_eval(dict_strings[2])
            return corrections, improvements, diction


        def get_development(essay, IeltsWritingScoreCritiria):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    请你先复习雅思写作评分标准："{IeltsWritingScoreCritiria}"
                    你是一个资深的雅思考官，你的任务是为一篇雅思大作文提供专业的修改建议和改进建议。请首先阅读这篇文章，然后生成一个string。
    
                    development = """
                    """
                    你需要针对文章的逻辑，发展和论证过程，论据是否充足4个方面指出它的问题以及如何提高：
                    
                    """
                    现在请你从阅读用户作文开始，用中文写。
                    我不需要任何警告和道歉。
                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    {essay}

                    '''
                }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo-16k",
                messages=messages,
                temperature=0,
                max_tokens=5000,
                #stream=True
            )
            return response.choices[0].message.content.strip()


        def get_rating(essay, essay_question, IeltsWritingScoreCritiria, Sampleanswers ):

            messages = [
                {
                    "role": "system", 
                    "content": 
                        f'''
                        你是一个严苛的资深雅思考官，你的任务是为一篇雅思大作文打分：

                        请你复习雅思写作评分标准："{IeltsWritingScoreCritiria}"
                        以及一篇6分范文："{Sampleanswers}"
                        然后为用户的文章打分，你的打分包括总分以及4项得分。永远是一个相差1分范围，比如1-2 分，或者5-6分。
                        记得，你很严苛，打分偏低。
                        然后写详细的点评，应该包括以下四个方面：

                        1. Task Response：评价文章是否明确给出了自己的观点，是否使用了足够的例子进行论证，是否准确理解了题目，是否有效地应对了任务。
                        2. Coherence and Cohesion：评价文章内容是否连贯流畅，段落之间是否使用了逻辑词语进行衔接，文章结构是否完整。
                        3. Lexical Resource：评价文章是否使用了丰富的词汇，是否使用了高级词组，词汇运用能力如何。
                        4. Grammatical Range and Accuracy：评价文章是否灵活运用了各种复杂语法结构，语法使用是否准确。
                        
                        Task Response、Coherence and Cohesion、Lexical Resource、Grammatical Range and Accuracy 保持英文，后面的点评用中文。
                        
                        接着，你要写一段中文，详细解释这篇文章在Task Response 上的逻辑问题以及如何提升
                        现在请你从阅读用户作文开始。我不需要任何警告和道歉。

                        '''
                },  
                {
                    "role": "user", 
                    "content": 
                        f'''
                        题目：{essay_question}\n\n
                        用户作文：{essay}            
                '''
                        
                }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=messages,
                temperature=0,
                max_tokens=5000,
                #stream=True
            )
            return response.choices[0].message.content.strip()


        def get_sample(essay_question, IeltsWritingScoreCritiria):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    你是一个资深的雅思考官，你的任务是写一篇8分的雅思大作文。主题与角度与原文章一致，写作水平要达到GRE满分写作水平。
                    请你复习雅思写作评分标准："{IeltsWritingScoreCritiria}"
                    然后从阅读题目开始。我不需要任何警告和道歉。

                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    {essay_question}
                    '''
                            }
                ]
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo-16k",
                messages=messages,
                temperature=0,
                max_tokens=5000,
                #stream=True
            )
            return response.choices[0].message.content.strip()


        def get_sample_rating(essay_question, sample, IeltsWritingScoreCritiria):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    你是一个资深的雅思考官，你的任务是写一篇8分的雅思大作文rating.
                    请你先复习雅思写作评分标准："{IeltsWritingScoreCritiria}"
                    你的评价应该包括以下四个方面：

                    1. Task Response：评价文章是否明确给出了自己的观点，是否使用了足够的例子进行论证，是否准确理解了题目，是否有效地应对了任务。
                    2. Coherence and Cohesion：评价文章内容是否连贯流畅，段落之间是否使用了逻辑词语进行衔接，文章结构是否完整。
                    3. Lexical Resource：评价文章是否使用了丰富的词汇，是否使用了高级词组，词汇运用能力如何。
                    4. Grammatical Range and Accuracy：评价文章是否灵活运用了各种复杂语法结构，语法使用是否准确。

                    Task Response、Coherence and Cohesion、Lexical Resource、Grammatical Range and Accuracy 保持英文，后面的点评用中文。
                    
                    现在请你从阅读题目和范文开始。我不需要任何警告和道歉。

                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    题目："{essay_question}"\n
                    范文："{sample}"\n
                            '''
                    }
                ]
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo-16k",
                messages=messages,
                temperature=0,
                max_tokens=5000,
                #stream=True
            )
            return response.choices[0].message.content.strip()

        def get_material(essay_question, sample):
            messages = [
                {
                    "role": "system", 
                    "content": 
                        f'''
                        你是一个资深的雅思考官，你的任务是为一篇雅思大作文提供写作素材。

                        你需要针对题目用中文写三个高中生可用的素材。
                        每个素材300中文字，你要非常详细地解释素材如何在文章中应用，包括可能使用到的高级单词和长句。
                        以下是一个示范：
                        [
                            素材一：极限运动对于压力的释放

                            极限运动与“是否应该禁止极限运动”这个题目的关联主要是，极限运动作为一种特殊的压力释放方式，为现代人提供了处理压力的有效途径。在今天的社会中，工作压力、学习压力以及生活压力常常使人们感到困扰，如果禁止了极限运动，那么这些寻找刺激和释放压力的人将失去一个有效的出口。例如，当人们参与跳伞或者滑雪这样的极限运动时，他们必须全神贯注地投入，这样就会暂时忘记日常生活中的琐碎小事，专注力的集中有助于他们从繁重的压力中暂时解脱出来，达到了压力释放的目的。因此，如果因为安全问题而全面禁止极限运动，可能会剥夺一些人处理压力的方式。在文章中，可以使用高级词汇如"relieve"（释放）、"pressure"（压力）、"diversity"（多样性）等。长句可以是："People can get access to sky diving to relieve since they are able to forget some trivial things in life when doing those things if they are permitted to do sky diving."

                            素材二：极限运动对于经济的推动

                            从经济的角度来看，极限运动与"是否应该禁止极限运动"这个主题有着密切的联系。极限运动，因其独特的刺激和挑战性，吸引了大量热衷于寻找新鲜刺激的人们。这些人通常愿意支付额外的费用去尝试这些运动，例如，他们可能会花费大量金钱去海底潜水，欣赏清澈的海洋和丰富多样的海洋生物，或者支付额外的费用去体验跳伞或冲浪的刺激。这种消费行为无疑为相关的旅游业和体育产业带来了巨大的经济收益。如果禁止了极限运动，那么可能会对相关的经济产业产生负面的影响，对于一些依赖于极限运动的地方经济，可能会造成相当大的打击。在文章中，可以使用高级词汇如"promote"（推动）、"economy"（经济）、"tourists"（游客）等。长句可以是："A large amount of people are willing to pay extra fees to dive under seas to expose themselves into crystal oceans appreciate a great diversity of sea creatures."

                            素材三：极限运动的禁止对于个人自由的限制

                            极限运动与"是否应该禁止极限运动"的主题之间的关系，可以从个人自由的角度进行理解。每个人都有选择自己生活方式的权利，包括选择自己喜欢的运动方式。极限运动虽然存在风险，但每个参与者都是自愿的，他们理解并接受这些风险，并且愿意承担可能出现的后果。如果政府因为安全问题而禁止极限运动，可能会侵犯到这些人的选择自由。人们有权选择他们认为对他们有益，能让他们感到快乐和满足的活动，无论这些活动是否具有风险。因此，从个人自由的角度看，禁止极限运动并不是一个好的选择。在文章中，可以使用高级词汇如"restrictions"（限制）、"freedom"（自由）、"rights"（权利）等。长句可以是："There is no need for government to ban on extreme sports since it would restrict people's freedom to choose their favorite sports."
                        ]
                        现在请你从阅读题目开始。我不需要任何警告和道歉。

                        '''
                },  
                {
                    "role": "user", 
                    "content": 
                        f'''
                        题目："{essay_question}"\n
                        8分范文："{sample}"\n          
                        '''
                }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo-16k",
                messages=messages,
                temperature=0,
                max_tokens=5000,
                #stream=True
            )
            return response.choices[0].message.content.strip()

        def modify_word_document(essay):
            # Load the document
            doc = Document()

            #get variables
            wordCount = count_words(essay)
            corrections, improvements, diction = get_correctionsandimprovements(essay)
            print('1')
            rating = get_rating(essay, essay_question, IeltsWritingScoreCritiria, Sampleanswers )
            development = get_development(essay, IeltsWritingScoreCritiria)
            print('2')
            sample = get_sample(essay_question, IeltsWritingScoreCritiria)
            sample_rating = get_sample_rating(essay_question, sample, IeltsWritingScoreCritiria)
            materials = get_material(essay_question, sample)
            print('3')
            # Add a paragraph with the essay
            doc.add_paragraph(essay)
            # Create a new document for the corrected text
            corrected_doc = Document()

            # Add a note at the beginning of the document
            # note_paragraph = corrected_doc.add_paragraph("雅思大作文批改")
            # run = note_paragraph.runs[0]
            # run.font.size = Pt(24)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7) 
            # note_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('雅思大作文批改', 0)


            note_paragraph = corrected_doc.add_paragraph(f'习作字数为{wordCount}')
            run = note_paragraph.runs[0]
            run.font.size = Pt(12)
            note_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            print(f"1")
            
            # Iterate over all paragraphs in the document
            for paragraph in doc.paragraphs:
                # Create a new paragraph with the corrected text
                new_paragraph = corrected_doc.add_paragraph()

                # Split the paragraph into sentences
                sentences = paragraph.text.split('.')

                # Iterate over all sentences in the paragraph
                for sentence in sentences:
                    # If the sentence is empty, skip it
                    if not sentence.strip():
                        continue

                    # Check if the sentence needs correction
                    for incorrect, correct in corrections.items():
                        if incorrect in sentence:
                            start, end = diff_strings(incorrect, correct)
                            incorrect_part = incorrect[start:end]
                            corrected_part = correct[start:end]

                            # Split the sentence around the incorrect part
                            before, _, after = sentence.partition(incorrect_part)

                            # Add the part before the incorrect part
                            new_paragraph.add_run(before)

                            # Add the incorrect part with strike-through
                            run = new_paragraph.add_run(incorrect_part)
                            run.font.strike = True

                            # Add the corrected part in brackets and color it red
                            corrected_run = new_paragraph.add_run(" (" + corrected_part + ")")
                            corrected_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # RGB for red

                            # Add the part after the incorrect part
                            new_paragraph.add_run(after + ".")

                            break  # Break the loop once a correction is applied

                    else:
                        # If the sentence doesn't need correction, add it without any modification
                        new_paragraph.add_run(sentence + "")

                
                # Add improvement suggestions after each paragraph
                for original, improvement in improvements.items():
                    if original in paragraph.text:

                        # Check if "建议改为" is in the improvement
                        if "建议改为" not in improvement:
                            # If not, prepend "建议改为"
                            improvement = "建议改为" + improvement

                        # Split on "建议改为" to get the suggestion
                        suggestion = improvement.split("建议改为")[1]

                        # Create a new paragraph for the suggestion
                        improved_paragraph = corrected_doc.add_paragraph()
                        
                        # Add the original sentence with underline
                        run = improved_paragraph.add_run(original + "")
                        run.font.underline = True

                        # Add "建议改为"
                        improved_paragraph.add_run("建议改为")

                        # Add the suggestion in blue
                        run = improved_paragraph.add_run(suggestion)
                        run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7) # RGB for blue

                # Add improvement suggestions after each paragraph
                for original, paraphrase in diction.items():
                    if original in paragraph.text:

                        # Check if "建议改为" is in the improvement
                        if "建议用" not in paraphrase:
                            # If not, prepend "建议改为"
                            paraphrase = "建议用" + paraphrase

                        # Split on "建议改为" to get the suggestion
                        suggestion = paraphrase.split("建议用")[1]

                        # Create a new paragraph for the suggestion
                        improved_paragraph = corrected_doc.add_paragraph()
                        
                        # Add the original sentence with underline
                        run = improved_paragraph.add_run(original + "")
                        run.font.underline = True

                        # Add "建议改为"
                        improved_paragraph.add_run("建议改为")

                        # Add the suggestion in blue
                        run = improved_paragraph.add_run(suggestion)
                        run.font.color.rgb = RGBColor(0x1C, 0x66, 0x36) # RGB for green
            print(f'1')
            # Add IELTS comments
            # Start a new page
            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)
            # title_line = corrected_doc.add_paragraph("文章评分:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('文章评分', 1)                          
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(rating)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Add IELTS comments

            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)
            # title_line = corrected_doc.add_paragraph("提分指南:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('提分指南', 1)  
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(development)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
            # Add a model essay

            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)
            # title_line = corrected_doc.add_paragraph("8分范文:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('8分范文', 1)  
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(sample)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Add an analysis of the model essay

            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)
            # title_line = corrected_doc.add_paragraph("8分范文点评:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('8分范文点评', 1)  
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(sample_rating)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
            # Add details for this essay question

            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)
            # title_line = corrected_doc.add_paragraph("素材指南:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('素材指南', 1) 
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(materials)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
                
            # Save the modified document
            filename = '/home/ec2-user/mysite/static/dlFile/IeltsTask2_modified.docx'
            corrected_doc.save(filename)

            # 返回相对路径
            return filename
        # Call your function that generates the file
        full_filepath = modify_word_document(essay)
        update_generate_count(-40)

        directory, filename = os.path.split(full_filepath)
        print(f"Trying to send file: {filename} from directory: {directory}")

        # Create a response with the file
        response = make_response(send_file(full_filepath, as_attachment=True))
        response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        return response

    except Exception as e:
        print(f"Error: {str(e)}")
        traceback.print_exc()
        return jsonify(status=500, message="服务器错误。", data={"error": str(e)})




#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#
#------------------ generate_toeflTaks1_report --------------------------#


@essayRating_blueprint.route("/api/generate_toeflTaks1_report", methods=["POST"])
@jwt_required()
@check_generate_count(40)
def generate_toeflTaks1_report():
    
    IntegratedWritingScoreCritiria = f'''
    TOEFL iBT® Integrated Writing Rubric

    Score Description:

    5 - A response at this level successfully selects the important information from the lecture and coherently and accurately presents this information in relation to the relevant information presented in the reading. The response is well organized, and occasional language errors that are present do not result in inaccurate or imprecise presentation of content or connections.

    4 - A response at this level is generally good in selecting the important information from the lecture and in coherently and accurately presenting this information in relation to the relevant information in the reading, but it may have minor omissions, inaccuracies, vagueness, or imprecision of some content from the lecture or in connection to points made in the reading. A response is also scored at this level if it has more frequent or noticeable minor language errors, as long as such usage and grammatical structures do not result in anything more than an occasional lapse of clarity or in the connection of ideas.

    3 - A response at this level contains some important information from the lecture and conveys some relevant connection to the reading, but it is marked by one or more of the following:

    Although the overall response is definitely oriented to the task, it conveys only vague, global, unclear, or somewhat imprecise connections of the points made in the lecture to points made in the reading.
    The response may omit one major key point made in the lecture.
    Some key points made in the lecture or the reading, or connections between the two, may be incomplete, inaccurate, or imprecise.
    Errors of usage and/or grammar may be more frequent or may result in noticeably vague expressions or obscured meanings in conveying ideas and connections.
    2 - A response at this level contains some relevant information from the lecture but is marked by significant language difficulties or by significant omission or inaccuracy of important ideas from the lecture or in the connections between the lecture and the reading. A response at this level is marked by one or more of the following:

    The response significantly misrepresents or completely omits the overall connection between the lecture and the reading.
    The response significantly omits or significantly misrepresents important points made in the lecture.
    The response contains language errors or expressions that largely obscure connections or meaning at key junctures or that would likely obscure understanding of key ideas for a reader not already familiar with the reading and the lecture.
    1 - A response at this level is marked by one or more of the following:

    The response provides little or no meaningful or relevant coherent content from the lecture.
    The language level of the response is so low that it is difficult to derive meaning.
    0 - A response at this level merely copies sentences from the reading, rejects the topic, or is otherwise not connected to the topic, is written in a foreign language, consists of keystroke characters, or is blank.
            '''
    SampleAnswer = f'''
    Writing (Integrated): Passage, Lecture, and Question
    Directions: Give yourself 3 minutes to read the passage.
    Reading Time: 3 minutes
    Critics say that current voting systems used in the United States are inefficient and often lead to
    the inaccurate counting of votes. Miscounts can be especially damaging if an election is closely
    contested. Those critics would like the traditional systems to be replaced with far more efficient
    and trustworthy computerized voting systems.
    In traditional voting, one major source of inaccuracy is that people accidentally vote for the
    wrong candidate. Voters usually have to find the name of their candidate on a large sheet of
    paper containing many names—the ballot—and make a small mark next to that name. People
    with poor eyesight can easily mark the wrong name. The computerized voting machines have an
    easy-to-use touch-screen technology: to cast a vote, a voter needs only to touch the candidate’s
    name on the screen to record a vote for that candidate; voters can even have the computer
    magnify the name for easier viewing.
    Another major problem with old voting systems is that they rely heavily on people to count the
    votes. Officials must often count up the votes one by one, going through every ballot and
    recording the vote. Since they have to deal with thousands of ballots, it is almost inevitable that
    they will make mistakes. If an error is detected, a long and expensive recount has to take place.
    In contrast, computerized systems remove the possibility of human error, since all the vote
    counting is done quickly and automatically by the computers.
    Finally some people say it is too risky to implement complicated voting technology nationwide.
    But without giving it a thought, governments and individuals alike trust other complex computer
    technology every day to be perfectly accurate in banking transactions as well as in the
    communication of highly sensitive information. \n\n 

    Directions: Here is the transcript.
    Narrator Now listen to part of a lecture on the topic you just read about.
    Professor While traditional voting systems have some problems, it’s doubtful that
    computerized voting will make the situation any better. Computerized voting may
    seem easy for people who are used to computers. But what about people who
    aren’t? People who can’t afford computers, people who don’t use them on a
    regular basis—these people will have trouble using computerized voting
    machines. These voters can easily cast the wrong vote or be discouraged from
    voting altogether because of fear of technology. Furthermore, it’s true that
    humans make mistakes when they count up ballots by hand. But are we sure that
    computers will do a better job? After all, computers are programmed by humans,
    so “human error” can show up in mistakes in their programs. And the errors
    caused by these defective programs may be far more serious. The worst a human
    official can do is miss a few ballots. But an error in a computer program can result
    in thousands of votes being miscounted or even permanently removed from the
    record. And in many voting systems, there is no physical record of the votes, so a
    computer recount in the case of a suspected error is impossible! As for our trust of
    computer technology for banking and communications, remember one thing:
    these systems are used daily and they are used heavily. They didn’t work
    flawlessly when they were first introduced. They had to be improved on and
    improved on until they got as reliable as they are today. But voting happens only
    once every two years nationally in the United States and not much more than
    twice a year in many local areas. This is hardly sufficient for us to develop
    confidence that computerized voting can be fully trusted.
    Directions: Give yourself 20 minutes to plan and write your response. Your response is judged
    on the quality of the writing and on how well it presents the points in the lecture and their
    relationship to the reading passage. Typically, an effective response will be 150 to 225 words.
    You may view the reading passage while you respond.
    Response time: 20 minutes
    Question: Summarize the points made in the lecture, being sure to explain how they cast doubt
    on specific points made in the reading passage. \n\n 

    Writing Practice Set 3 (Integrated): Sample Responses
    Response A, Score of 5
    The lecture explained why the computerized voting system can not replace the traditional voting
    system. There are the following three reasons.
    First of all, not everyoen one can use computers correctly. Some people do not have access to
    computers, some people are not used of computers, and some people are even scared of this new
    technology. If the voters do not know how to use a computer, how do you expect them to finish
    the voting process through computers? This directly refutes the reading passage which states that
    computerized voting is easier by just touchingthe screen.
    Secondly, computers may make mistakes as the people do. As computers are programmed by the
    human beings, thus erros are inevitable in the computer system. Problems caused by computer
    voting systems may be more serious than those caused by people. A larger number of votes
    might be miss counted or even removed from the system. Furthermore, it would take more
    energy to recount the votes. Again this contradicts what is stated in the reading which stated that
    only people will make mistakes in counting.
    Thirdly, computerized voting system is not reliable because it has not reached a stable status.
    People trust computers to conduct banking transactions because the computerized banking
    system is being used daily and frecuently and has been stable. How ever, the voting does not
    happen as often as banking thus the computerized voting system has not been proved to be
    totally reliable.
    All in all, not everyone can use a computer properly, computer cause mistakes and computerized
    voting system is not reliable are the main reasons why computerized voting system can not
    replace the traditional voting system.
    Score explanation
    This response is well organized, selects the important information from all three points made in
    the lecture, and explains its relationship to the claims made in the reading passage about the
    advantages of computerized voting over traditional voting methods.
    First, it counters the argument that computerized voting is more user-friendly and prevents
    distortion of the vote by saying that many voters find computers unfamiliar and some voters may
    end up not voting at all.
    Second, it challenges the argument that computerized voting will result in fewer miscounts by
    pointing out that programming errors may result in large-scale miscounts and that some errors
    may result in the loss of voting records.
    Third, it rejects the comparison of computerized voting with computerized banking by pointing
    out that the reliability of computerized banking (“reached a stable status”) has been achieved
    though frequent use, which does not apply to voting. 
    There are occasional minor language errors: for example, “people not used of computers”; “miss
    counted”; “computer cause mistakes”; and the poor syntax of the last sentence (“All in all . . . ”).
    Some spelling errors are obviously typos: “everyoen.” The errors, however, are not at all
    frequent and do not result in unclear or inaccurate representation of the content.
    The response meets all the criteria for the score of 5 \n 

    Response B, Score of 4
    The leture disgreed with the article's opinions. It's not a better solution to use the computerized
    voting systems.
    Firstly, it might be hard for the voters who don't use the computer so often, or the users who is
    fear of the technology, even some of voters can not aford a computer. Touch screen may also be
    hard to use for people who is not familiar with computers. Secondly, computer is programmed
    by human beings, which means it can also have errors. Instead of human being's counting error,
    which only results one or two counting error in number, an errror in the program code could
    cause tramendous error in number. In case of the computer crash or disaster, it may lost all the
    voting information. We can not even to make a re-count. Lastly, our daily banking or other
    highly sensitive infomation system, is actually improved as time goes by. They were also
    problematic at the beginning. As we use them so often, we have more chances to find problems,
    and furturemore, to fix and improve them. However, for the voting system, we only use them
    every 2 years nationally and some other rare events. We just don't use it often enough to find a
    bug or test it thoroughly.
    Score explanation
    The response selects most of the important information from the lecture and indicates that it
    challenges the main argument in the reading passage about the advantages of computerized
    voting systems (“it’s not a better solution”).
    First, the response explains that some people will not find computers to be user- friendly;
    however, it fails to relate this clearly to the point made in the passage that computerized voting
    will prevent distortion of the vote. That is clearly an omission, but it is minor.
    Second, the response does a good job of pointing out how programming and errors can cause
    greater problems than miscounts cause in the traditional voting system.
    Third, the response provides a nice explanation of how the frequent use of systems like the
    banking system has contributed to such systems’ reliability, and then it contrasts that with the
    computerized voting system.
    There are more frequent language errors throughout the response—for example, “users who is
    fear”; “some of voters can not aford”; “people who is not familiar”; “it may lost”; and “can not
    even to make.” Expressions chosen by the writer occasionally affect the clarity of the content
    that is being conveyed: “results one or two counting error in number . . . an errror in the program
    code could cause tramendous error in number” and “use them every 2 years nationally and some 
    Copyright © 2021 by Educational Testing Service. All rights reserved. ETS, the ETS logo, TOEFL and TOEFL iBT are registered trademarks of Educational Testing
    Service (ETS) in the United States and other countries. IN ENGLISH WITH CONFIDENCE is a trademark of ETS.
    other rare events.” However, it should be noted that in these cases, a reader can derive the
    intended meaning from the context.
    Due to the more frequent language errors that on occasion result in minor lapses of clarity and
    due to minor content omission, especially in the coverage of the first lecture point, the response
    cannot earn the score of 5. At the same time, since the language errors are generally minor and
    mostly do not interfere with the clarity of the content and since most of the important
    information from the lecture is covered by the writer, the response deserves a higher score than
    3. It meets the criteria for the score of 4.\n 
    '''
    
    
    try:
        data = request.get_json()
        prompt = data.get("essay_question")
        essay = data.get("essay")
        essay_question = prompt + "Read listening script of a short lecture and the short passage, then write in response to what you read and listened to."
        # 加载敏感词
        sensitive_words = load_sensitive_words('chinese')

        # 检查是否包含敏感词
        contains_sensitive = contains_sensitive_words(essay_question, sensitive_words)
        if contains_sensitive:
            return jsonify({"error": "您的输入包含敏感词，请重新输入。"}), 400
        
        #---------------------------Def--------------------#
        def get_correctionsandimprovements(essay):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    你是一个资深的托福考官及python专家，你的任务是为一篇托福Integrated essay 提供专业的修改建议和改进建议。请首先阅读这篇文章，然后生成两个词典：

                    1. `corrections`：这是一个字典，键是文章中需要修正的部分，修改词汇用错、语法、标点符号用错的地方，值是你用中文写的建议的修正和修正理由。
                    2. `improvements`：这是一个字典，键是文章中可以改进的部分主要集中于Coherence & Cohesion, diction 和 sentence variety，值是你用中文写的建议的改进和改进理由。
                
                    两个字典都要列出3项。
                    
                    
                    你要写纯粹的字典格式，并且不包括任何其他说明或者引用符号。格式和下面的例子保持100%一致：
                    比如原文有一句“The speed of the car is high”，那么corrections和improvements修改格式如下：
                    corrections = {
                        "is fast" ':' "is high”在此处使用不当，建议改为The speed of the car is high”。因为..."
                    }
                    improvements = {
                        "The speed of the car is fast." ':' "建议改为The speed of the car is high,因为"
                    }
                    
                    """
                    现在请你从阅读用户作文开始。我不需要任何警告和道歉。

                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    {essay}
                    
                    '''
            }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            string = response.choices[0].message.content.strip()
                # 删除 Python 关键字
            response_string = string.replace("corrections = ", "").replace("improvements = ", "")
            # 将字符串分割成两部分，每部分包含一个字典
            dict_strings = response_string.split("\n\n")
            # 对每一部分使用 ast.literal_eval() 函数，得到两个字典
            corrections = ast.literal_eval(dict_strings[0])
            improvements = ast.literal_eval(dict_strings[1])
            
            return corrections, improvements

        def get_rating(essay, essay_question, IntegratedWritingScoreCritiria, SampleAnswer):
            
            messages = [
                {
                    "role": "system", 
                    "content": 
                        f'''
                        你是一个资深的托福考官，
                        请你复习这篇作文的评分标准：\n"{IntegratedWritingScoreCritiria}"
                        以及两篇范文及官方点评：\n"{SampleAnswer}"

                        你的评分非常严苛，一般略为偏低。
                        现在你的任务是为一篇托福作文打分, 然后写详细的点评。你需要明确指出学生习作与提供的听力材料和阅读文章之间的匹配度问题，
                        比如学生习作在概括听力材料的某些观点时,没有充分抓住听力材料的关键信息。
                        请你参考下面的建议：
                        【

                        Here is a summary of the key points for a high-scoring TOEFL Integrated Writing response:
                        Coherent organization and complete content coverage. A top-scoring response should have an effective structure to connect the three lecture points smoothly and cover the main information from the lecture.
                        Citing specific details and examples from the lecture, such as some voters' unfamiliarity with computer technology.
                        Clarifying relationships between opposing claims in the passage and lecture, like the difference in reliability of computerized voting versus banking systems.
                        Accurate and fluent language with minimal errors that do not affect meaning.
                        Appropriate word count, typically 150-225 words.
                        Logical paragraph divisions, such as presenting the first lecture point against the passage in one paragraph.
                        Clear logical flow with transitions between paragraphs and points.
                        No lapses in clarity or accuracy.
                        
                        】
                        你需要写英文和中文版。
                        现在请你从阅读用户作文开始，用中文写，以第二人称“你”点评，不需要有“English Version/Chinese Version“的字眼，也不需要任何警告和道歉。

                        '''
                },  
                {
                    "role": "user", 
                    "content": 
                        f'''
                        题目：{essay_question}\n\n
                        用户作文：{essay}            
                '''
                        
                }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            return response.choices[0].message.content.strip()

        def get_sample(essay_question, IntegratedWritingScoreCritiria, SampleAnswer):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    你是一个资深的托福考官，你的任务是写一篇满分托福的Integrated Task的作文。
                    请你复习这篇作文的评分标准："{IntegratedWritingScoreCritiria}"
                    以及两篇范文："{SampleAnswer}"
                    
                    你还可以参考以下的满分学生的建议：
                    [
                    Here is a summary of the key points for a high-scoring TOEFL Integrated Writing response:
                    Coherent organization and complete content coverage. A top-scoring response should have an effective structure to connect the three lecture points smoothly and cover the main information from the lecture.
                    Citing specific details and examples from the lecture, such as some voters' unfamiliarity with computer technology.
                    Clarifying relationships between opposing claims in the passage and lecture, like the difference in reliability of computerized voting versus banking systems.
                    Accurate and fluent language with minimal errors that do not affect meaning.
                    Appropriate word count, typically 150-225 words.
                    Logical paragraph divisions, such as presenting the first lecture point against the passage in one paragraph.
                    Clear logical flow with transitions between paragraphs and points.
                    No lapses in clarity or accuracy.
                    ]
                    
                    你要非常重视字数限制，不能超过200英文单词。
                    

                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    这是题目、阅读材料和听力材料：{essay_question}\n\n
                    请用英文写。
                    托福作文不需要title。
                    '''
            }
                ]
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo-16k",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            return response.choices[0].message.content.strip()

        def get_sample_rating(sample, IntegratedWritingScoreCritiria, SampleAnswer):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    你是一个资深的托福考官，你的任务是为一篇托福作文打分：

                    请你复习这篇作文的评分标准：\n"{IntegratedWritingScoreCritiria}"
                    以及两篇范文及官方点评：\n"{SampleAnswer}"

                    请写详细的点评但不要打分。
                    你应该模仿上面发给你的官方点评的风格，也可以参考下面的建议：
                    【
                        Here is a summary of the key characteristics of official scoring rubrics for the TOEFL Integrated Writing task:

                        Citing specific examples from the response
                        Analyzing and comparing treatment of lecture and passage points
                        Clearly stating which criteria are met or not met
                        Classifying types of errors and their effect on expression
                        Providing objective feedback
                        Having clear structure and accurate language
                        Giving constructive feedback for improvement
                        Using an appropriate word count (130-150 words)
                        In conclusion, official TOEFL writing rubrics have the characteristics of using examples, analysis, clear criteria, error classification, objectivity, structure, constructive feedback, and appropriate length. These features help students understand scoring standards and improve their writing skills.
                        
                        Relevant and elaboration, Effective use of language和Lexical or grammatical errors 保持英文，后面的点评用中文。
                    】
                        
                        现在请你从阅读用户作文开始，用中文写，以第二人称“你”点评，你的评分一定是5分，也不需要任何警告和道歉。
                    

                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    题目："{essay_question}"\n
                    这是一篇满分范文，请解析它的优点："{sample}"\n
                            '''
                    }
                ]
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo-16k",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            return response.choices[0].message.content.strip()

        def get_ideas(essay_question, IntegratedWritingScoreCritiria, SampleAnswer):
            print(f'{essay_question}')
            messages = [
                {
                    "role": "system", 
                    "content": 
                        f'''
                        你是一个资深的托福考官，你的任务是为一篇托福作文提供不同的写作思路：

                        请你复习这篇作文的评分标准：\n"{IntegratedWritingScoreCritiria}"
                        以及两篇范文及官方点评：\n"{SampleAnswer}"

                        请你展开详细叙述关于这篇文章的建议，每一点都要针对本题目举例,给出一些可能用到的地道的表达方式：
                        【                       
                        注意文章与讲座的观点对比
                        
                        概括每段讲座内容与文章的关系
                        
                        使用恰当的转折词(如however, in contrast...)
                        
                        适当引用文章和讲座细节
                        
                        
                        】
                        现在请你从阅读用户作文开始，用中文写，以第二人称“你”点评，也不需要任何警告和道歉。

                        '''
                },  
                {
                    "role": "user", 
                    "content": 
                        f'''
                        题目："{essay_question}"\n
        
                        '''
                }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo-16k",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            return response.choices[0].message.content.strip()

        def modify_word_document(essay):
            # Load the document
            doc = Document()

            wordCount = count_words(essay)
            corrections, improvements = get_correctionsandimprovements(essay)
            print('1')
            rating = get_rating(essay, essay_question, IntegratedWritingScoreCritiria, SampleAnswer )
            ideas = (essay_question, IntegratedWritingScoreCritiria, SampleAnswer)
            print('2')
            sample = get_sample(essay_question, IntegratedWritingScoreCritiria, SampleAnswer)
            sample_rating = get_sample_rating(sample, IntegratedWritingScoreCritiria, SampleAnswer)
            ideas = get_ideas(essay_question, IntegratedWritingScoreCritiria, SampleAnswer)
            print('3')  
            # Add a paragraph with the essay
            doc.add_paragraph(essay)
            # Create a new document for the corrected text
            corrected_doc = Document()
            

            # Add a note at the beginning of the document
            # note_paragraph = corrected_doc.add_paragraph("托福综合作文批改")
            # run = note_paragraph.runs[0]
            # run.font.size = Pt(24)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7) 
            # note_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('托福综合作文批改', 0) 

            note_paragraph = corrected_doc.add_paragraph(f'习作字数为{wordCount}')
            run = note_paragraph.runs[0]
            run.font.size = Pt(12)
            note_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Iterate over all paragraphs in the document
            for paragraph in doc.paragraphs:

                # Create a new paragraph with the corrected text
                new_paragraph = corrected_doc.add_paragraph()

                # Split the paragraph into sentences
                sentences = paragraph.text.split('.')

                # Iterate over all sentences in the paragraph
                for sentence in sentences:
                    # If the sentence is empty, skip it
                    if not sentence.strip():
                        continue

                    # Check if the sentence needs correction
                    for incorrect, correct in corrections.items():
                        if incorrect in sentence:
                            start, end = diff_strings(incorrect, correct)
                            incorrect_part = incorrect[start:end]
                            corrected_part = correct[start:end]

                            # Split the sentence around the incorrect part
                            before, _, after = sentence.partition(incorrect_part)

                            # Add the part before the incorrect part
                            new_paragraph.add_run(before)

                            # Add the incorrect part with strike-through
                            run = new_paragraph.add_run(incorrect_part)
                            run.font.strike = True

                            # Add the corrected part in brackets and color it red
                            corrected_run = new_paragraph.add_run(" (" + corrected_part + ")")
                            corrected_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # RGB for red
                            corrected_run.font.name = 'Microsoft YaHei'
                            # Add the part after the incorrect part
                            new_paragraph.add_run(after + ".")

                            break  # Break the loop once a correction is applied

                    else:
                        # If the sentence doesn't need correction, add it without any modification
                        new_paragraph.add_run(sentence + ".")

                
                # Add improvement suggestions after each paragraph
                for original, improvement in improvements.items():
                    if original in paragraph.text:

                        # Check if "建议改为" is in the improvement
                        if "建议改为" not in improvement:
                            # If not, prepend "建议改为"
                            improvement = "建议改为" + improvement

                        # Split on "建议改为" to get the suggestion
                        suggestion = improvement.split("建议改为")[1]

                        # Create a new paragraph for the suggestion
                        improved_paragraph = corrected_doc.add_paragraph()
                        
                        # Add the original sentence with underline
                        run = improved_paragraph.add_run(original + ". ")
                        run.font.underline = True
                        run.font.size = Pt(12)
                        run.font.name = 'Microsoft YaHei'
                        # Add "建议改为"
                        improved_paragraph.add_run("建议改为")

                        # Add the suggestion in blue
                        run = improved_paragraph.add_run(suggestion)
                        run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)  # RGB for blue

            # Add essay comments
            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)   
            # title_line = corrected_doc.add_paragraph("评分:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.name = 'Microsoft YaHei'
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('评分', 1) 
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(rating)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(12)
            

            # Add details for this essay question
            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)   
            # title_line = corrected_doc.add_paragraph("思路建议:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('思路建议', 1) 
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(ideas)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(12)
            

            # Add a model essay
            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)   
            # title_line = corrected_doc.add_paragraph("满分范文:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('满分范文', 1) 
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(sample)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            

            # Add an analysis of the model essay
            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)   
            # title_line = corrected_doc.add_paragraph("满分范文点评:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('满分范文点评', 1) 
            # Add the content with Times New Roman font
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(sample_rating)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(12)
                    
            # Save the modified document
            filename = '/home/ec2-user/mysite/static/dlFile/ToeflTask1_modified.docx'
            corrected_doc.save(filename)

            # 返回相对路径
            return filename
        # Call your function that generates the file
        full_filepath = modify_word_document(essay)
        update_generate_count(-40)

        directory, filename = os.path.split(full_filepath)
        print(f"Trying to send file: {filename} from directory: {directory}")

        # Create a response with the file
        response = make_response(send_file(full_filepath, as_attachment=True))
        response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        return response

    except Exception as e:
        print(f"Error: {str(e)}")
        traceback.print_exc()
        return jsonify(status=500, message="服务器错误。", data={"error": str(e)})
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#
#------------------ generate_toeflTaks2_report --------------------------#


@essayRating_blueprint.route("/api/generate_toeflTaks2_report", methods=["POST"])
@jwt_required()
@check_generate_count(40)
def generate_toeflTaks2_report():
    
    AcademicDiscussionScoreCritiria = f'''
    Writing for an Academic Discussion Rubric:

    Score 5:
    A fully successful response

    The response is highly relevant and very clearly expressed, making a valuable contribution to the online discussion.
    Demonstrates consistent proficiency in language use.
    Typical response characteristics:
    Provides relevant and well-elaborated explanations, exemplifications, and/or details.
    Exhibits effective use of a variety of syntactic structures and precise, idiomatic word choice.
    Almost no lexical or grammatical errors, except for minor mistakes common in competent writing under timed conditions (e.g., typos, common misspellings, or substitutions like there/their).
    \n\n
    Score 4:
    A generally successful response

    The response is relevant and contributes effectively to the online discussion.
    The writer's ideas are easily understood due to their proficiency in language use.
    Typical response characteristics:
    Offers relevant and adequately elaborated explanations, exemplifications, and/or details.
    Utilizes a variety of syntactic structures and appropriate word choice.
    Few lexical or grammatical errors.
    \n\n
    Score 3:
    A partially successful response

    The response is mostly relevant and understandable, but there might be some limitations in language use.
    Typical response characteristics:
    Contains elaborations with some parts of explanations, examples, or details that are missing, unclear, or irrelevant.
    Demonstrates some variety in syntactic structures and uses a range of vocabulary.
    Notable lexical and grammatical errors in sentence structure, word form, or use of idiomatic language may be present.
    \n\n
    \n\n
    Score 2:
    A mostly unsuccessful response

    The response reflects an attempt to contribute to the online discussion, but language limitations may make ideas hard to follow.
    Typical response characteristics:
    Ideas may be poorly elaborated or only partially relevant.
    Shows a limited range of syntactic structures and vocabulary.
    Accumulation of errors in sentence structure, word forms, or use may hinder comprehension.
    \n\n
    Score 1:
    An unsuccessful response

    The response reflects an ineffective attempt to contribute to the online discussion due to severe language limitations.
    Typical response characteristics:
    Contains words and phrases that indicate an attempt to address the task but with few or no coherent ideas.
    Demonstrates a severely limited range of syntactic structures and vocabulary.
    Serious and frequent errors in language use are evident.
    Minimal original language; any coherent language is mostly borrowed from the stimulus.

            '''
    SampleAnswer = f'''
    Your professor is teaching a class on economics. Write a post responding to the professor’s
    question.
    In your response you should:
    • express and support your opinion
    • make a contribution to the discussion
    An effective response will contain at least 100 words. You will have 10 minutes to write it.
    Dr. Achebe
    When people are asked about the most important discoveries or inventions made
    in the last two hundred years, they usually mention something very obvious, like
    the computer or the cell phone. But there are thousands of other discoveries or
    inventions that have had a huge impact on how we live today. What scientific
    discovery or technological invention from the last two hundred years—other than
    computers and cell phones—would you choose as being important? Why?
    Paul
    I mean, we’re so used to science and technology that we are not even aware of all
    the things we use in our daily lives. I would probably choose space satellites. This
    technology happened in the last hundred years, and it has become important for so
    many things. Just think about navigation, or telecommunications, or even the
    military.
    Claire
    I am thinking about medical progress. Like, for example, when scientists
    discovered things about healthy nutrition. I am thinking of identifying all the
    vitamins we need to stay healthy. I am not sure exactly when the vitamin
    discoveries happened, but I know they are very important. Our health is much
    better than it was 200 years ago.
    \n\n
        Response A, Score of 5
    In the past 200 years, tons of scientific discoveries or technological inventions have been shown
    to the world. If I had to choose one in particular it will probably be vaccine or antibiotics. With
    Pasteur's work and discoveries, the world changed in a way people couldn't imagine. So many
    people were dying really young because at that time life's conditions were not as good as the one
    we have now. With vaccine, we could now irradicate diseases that were killing millions of
    people, we learn so much about the immune system and ways our body was reacting to
    pathogens and the answers he could produce to defend us against it. Medicine evolved so much
    and keeps evolving every day because scientists are curious to understand how our body is
    working and how he is able to communicate with our environment. People aged 40 are now not
    that old and still have a really long life to live and enjoy when 2 centuries ago it was
    synonymous of 80% chance of dying.\n
    Score explanation\n
    This is a fully successful response. The writer chooses vaccines/antibiotics as the most important
    invention of the past 200 years. The author then provides a description, for contrast, of what life
    was like before Pasteur’s work (people dying young) and after the vaccine was created (millions
    of lives saved, more understanding of the immune system). The writer continues on to point out
    that medicine continues to evolve because of Pasteur’s work and how human lifespans have been
    extended. Overall, the response provides well-elaborated explanations and details to support the
    main opinion and provides a relevant contribution to the discussion.
    While there are almost no errors in grammar and word choice, there are a few minor ones that
    have little impact on meaning (such as “life's conditions were not as good as the one we have”
    rather than “were not as good as they are now,” and “how our body is working and how he is
    able to communicate” rather than “how it is able to communicate”). However, these errors are
    fairly minor and such errors might be expected when writing under timed conditions. The writer
    is able to use some complex sentences and relatively precise vocabulary, which is expected in a
    5-level response.\n\n

    Response B, Score of 4\n
    From my personal point of view, I think the most important invention is the light bulb. Before it
    was invented, people had have to use candles for illumination in the evening. It's performance is
    not very stable, and it is produce really high tempreture which would probably lead to a fire
    accident. Light bulbs, however, produce constant and bright lighting at nights. One light bulb
    could use for several years, which is quite convenient-people don't need to storage many bulbs.
    What's more, it is safer than past candles. This is a huge progress in technology, and I consider it
    as the most vital invention from the last 200 years.\n
    Score explanation\n
    This is a generally successful response. The writer chooses the light bulb as the most important
    invention and contrasts it with the disadvantages of using candles (their unstable performance,
    high temperature, and possibility of fire. The writer then names several advantages of the light
    bulb (a stable, long-lasting performance; brightness; and safety. Notice these points are in direct
    contrast to the mentioned drawbacks of candles, which is an effective way to support the choice.
    However, the argument might have been a little stronger if the writer had discussed not only why
    light bulbs are better than candles, but also why light bulbs are the most important invention of
    the past 200 years (other than computers and cell phones, perhaps by mentioning at least one
    way in which they have helped societies progress.
    While the writer uses a nice variety of grammatical constructions and fairly precise vocabulary,
    the number and type of errors in grammar and vocabulary prevent it from reaching the highest
    score level. Multiple small errors such as “had have to use”, “it is produce”, “one light bulb
    could use for years”, and “don’t need to storage” are distracting for the reader even though the
    intended meaning is still usually clear. 

    '''
    
    try:
        data = request.get_json()
        essay_question = data.get("essay_question")
        essay = data.get("essay")
        
        # 加载敏感词
        sensitive_words = load_sensitive_words('chinese')

        # 检查是否包含敏感词
        contains_sensitive = contains_sensitive_words(essay_question, sensitive_words)
        if contains_sensitive:
            return jsonify({"error": "您的输入包含敏感词，请重新输入。"}), 400

        def get_correctionsandimprovements(essay):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    你是一个资深的托福考官及python专家，你的任务是为一篇托福Academic Discussion essay 提供专业的修改建议和改进建议。请首先阅读这篇文章，然后生成两个词典：

                    1. `corrections`：这是一个字典，键是文章中需要修正的部分，修改词汇用错、语法、标点符号用错的地方，值是你用中文写的建议的修正和修正理由。
                    2. `improvements`：这是一个字典，键是文章中可以改进的部分主要集中于Coherence & Cohesion, diction 和 sentence variety，值是你用中文写的建议的改进和改进理由。
                
                    两个字典都要列出6项。
                    
                    
                    你要写纯粹的字典格式，并且不包括任何其他说明或者引用符号。格式和下面的例子保持100%一致：
                    比如原文有一句“The speed of the car is high”，那么corrections和improvements修改格式如下：
                    corrections = {
                        "is fast" ':' "is high”在此处使用不当，建议改为The speed of the car is high”。因为..."
                    }
                    improvements = {
                        "The speed of the car is fast." ':' "建议改为The speed of the car is high,因为"
                    }
                    
                    """
                    现在请你从阅读用户作文开始。Don't give me any warnings or apologies.

                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    {essay}
                    
                    '''
            }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            string = response.choices[0].message.content.strip()
                # 删除 Python 关键字
            response_string = string.replace("corrections = ", "").replace("improvements = ", "")
            # 将字符串分割成两部分，每部分包含一个字典
            dict_strings = response_string.split("\n\n")
            # 对每一部分使用 ast.literal_eval() 函数，得到两个字典
            corrections = ast.literal_eval(dict_strings[0])
            improvements = ast.literal_eval(dict_strings[1])
            
            return corrections, improvements

        def get_rating(essay, essay_question, AcademicDiscussionScoreCritiria, SampleAnswer):

            messages = [
                {
                    "role": "system", 
                    "content": 
                        f'''
                        你是一个资深的托福考官，你的任务是为一篇托福作文打分：

                        请你复习这篇作文的评分标准：\n"{AcademicDiscussionScoreCritiria}"
                        以及两篇范文及官方点评：\n"{SampleAnswer}"
                        首先请你为用户的文章打分,
                        然后写详细的点评，应该包括以下三方面：

                        1. Relevant and elaboration：
                        2. Effective use of language：
                        3. Lexical or grammatical errors：
                        
                        Relevant and elaboration, Effective use of language和Lexical or grammatical errors 保持英文，后面的点评用中文。
                        
                        现在请你从阅读用户作文开始，用中文写，以第二人称“你”点评，不需要有“English Version/Chinese Version“的字眼。Don't give me any warnings or apologies.

                        '''
                },  
                {
                    "role": "user", 
                    "content": 
                        f'''
                        题目：{essay_question}\n\n
                        用户作文：{essay}            
                '''
                        
                }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            return response.choices[0].message.content.strip()


        def get_development(essay, SampleAnswer):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    你是一个资深的托福考官，你的任务是为一篇托福Academic Discussion提供建议。
                    请首先参考两篇范文及官方点评：\n"{SampleAnswer}"
                    
                    然后阅读学生文章，然后生成一个字符串development。
                    development = """
                    """
                    你可以参考以下的满分学生的建议：
                    [
                    Meaningful Contribution to the Discussion: The essay needs to make a meaningful contribution to the ongoing discussion. This means offering an independent, insightful perspective, rather than simply repeating or superficially responding to others' points.
                    Clear Stance and Support: You need to articulate a clear stance on the discussion topic, and provide substantial support for your view. This could include evidence, examples, logical reasoning, or detailed explanations.
                    Precision and Fluency in Language Use: Your essay should demonstrate precision and fluency in your use of English. This includes correct grammar, spelling, and punctuation, as well as a variety of sentence structures and vocabulary.
                    Adapting to the Discussion Context and Tone: This is an online academic discussion, so you need to write in a language and style that suits this setting. While the language can be slightly informal, it needs to follow basic grammar rules, and disagreements should be expressed respectfully.
                    Adequate Development and Organization: Your essay should adequately develop your viewpoint and organize your thoughts in a clear way. This includes using appropriate paragraph breaks, making your stance and support easy to understand.
                    In summary, a high-scoring Academic Discussion essay needs to show in-depth thinking, clear expression of viewpoints, precision and fluency in language use, and good essay structure and organization.
                    ]
                    现在请你从阅读用户作文开始，用中文写。Don't give me any warnings or apologies.

                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    {essay}
                
                    '''
                }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            return response.choices[0].message.content.strip()


        def get_sample(essay_question, AcademicDiscussionScoreCritiria, SampleAnswer):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    你是一个资深的托福考官，你的任务是写一篇满分的Academic Discussion大作文。你的字数必须控制在200字以内。
                    请你复习这篇作文的评分标准："{AcademicDiscussionScoreCritiria}"
                    以及两篇范文："{SampleAnswer}"
                    
                    development = """
                    """
                    你可以参考以下的满分学生的建议：
                    [
                    Meaningful Contribution to the Discussion: The essay needs to make a meaningful contribution to the ongoing discussion. This means offering an independent, insightful perspective, rather than simply repeating or superficially responding to others' points.
                    Clear Stance and Support: You need to articulate a clear stance on the discussion topic, and provide substantial support for your view. This could include evidence, examples, logical reasoning, or detailed explanations.
                    Precision and Fluency in Language Use: Your essay should demonstrate precision and fluency in your use of English. This includes correct grammar, spelling, and punctuation, as well as a variety of sentence structures and vocabulary.
                    Adapting to the Discussion Context and Tone: This is an online academic discussion, so you need to write in a language and style that suits this setting. While the language can be slightly informal, it needs to follow basic grammar rules, and disagreements should be expressed respectfully.
                    Adequate Development and Organization: Your essay should adequately develop your viewpoint and organize your thoughts in a clear way. This includes using appropriate paragraph breaks, making your stance and support easy to understand.
                    In summary, a high-scoring Academic Discussion essay needs to show in-depth thinking, clear expression of viewpoints, precision and fluency in language use, and good essay structure and organization.
                    ]
                    现在请你从阅读题目开始。Don't give me any warnings or apologies.

                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    题目：{essay_question}\n\n
                    '''
            }
                ]
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            return response.choices[0].message.content.strip()

        def get_sample_rating(essay_question, sample, AcademicDiscussionScoreCritiria, SampleAnswer):
            messages = [
            {
                "role": "system", 
                "content": 
                    f'''
                    你是一个资深的托福考官，你的任务是为一篇托福作文打分：

                    请你复习这篇作文的评分标准：\n"{AcademicDiscussionScoreCritiria}"
                    以及两篇范文及官方点评：\n"{SampleAnswer}"

                    请你写详细的点评，你应该模仿官方点评的风格，且应该包括以下三方面：

                    1. Relevant and elaboration：
                    2. Effective use of language：
                    3. Lexical or grammatical errors：
                    
                    你需要写中文。
                    现在请你从阅读作文开始。Don't give me any warnings or apologies.
                    

                    '''
            },  
            {
                "role": "user", 
                "content": 
                    f'''
                    题目："{essay_question}"\n
                    范文："{sample}"\n
                            '''
                    }
                ]
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            return response.choices[0].message.content.strip()

        def get_ideas(essay_question, AcademicDiscussionScoreCritiria, SampleAnswer):
            messages = [
                {
                    "role": "system", 
                    "content": 
                        f'''
                        你是一个资深的托福考官，你的任务是为一篇托福作文提供不同的写作思路：

                        请你复习这篇作文的评分标准：\n"{AcademicDiscussionScoreCritiria}"
                        以及两篇范文及官方点评：\n"{SampleAnswer}"

                        请你展开详细叙述三种不同的写作思路：

                        现在请你从阅读作文开始。Don't give me any warnings or apologies.

                        '''
                },  
                {
                    "role": "user", 
                    "content": 
                        f'''
                        题目："{essay_question}"\n
        
                        '''
                }
            ]
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=messages,
                temperature=0,
                max_tokens=3000,
                #stream=True
            )
            return response.choices[0].message.content.strip()

        def modify_word_document(essay):
            # Load the document
            doc = Document()
            wordCount = count_words(essay)
            corrections, improvements = get_correctionsandimprovements(essay)
            print('1')
            rating = get_rating(essay, essay_question, AcademicDiscussionScoreCritiria, SampleAnswer)
            development = get_development(essay, SampleAnswer)
            print('2')
            sample = get_sample(essay_question, AcademicDiscussionScoreCritiria, SampleAnswer)
            sample_rating = get_sample_rating(essay_question, sample, AcademicDiscussionScoreCritiria, SampleAnswer)
            ideas = get_ideas(essay_question, AcademicDiscussionScoreCritiria, SampleAnswer)
            print('3')

            # Add the input text as a paragraph
            doc.add_paragraph(essay)
            # Create a new document for the corrected text
            corrected_doc = Document()

        
            # Add a note at the beginning of the document
            # note_paragraph = corrected_doc.add_paragraph("托福学术讨论作文批改")
            # run = note_paragraph.runs[0]
            # run.font.size = Pt(24)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7) 
            # note_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('托福学术讨论作文批改', 0) 
            

            note_paragraph = corrected_doc.add_paragraph(f'习作字数为{wordCount}')
            run = note_paragraph.runs[0]
            run.font.size = Pt(12)
            note_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Iterate over all paragraphs in the document
            for paragraph in doc.paragraphs:
                # Create a new paragraph with the corrected text
                new_paragraph = corrected_doc.add_paragraph()

                # Split the paragraph into sentences
                sentences = paragraph.text.split('.')

                # Iterate over all sentences in the paragraph
                for sentence in sentences:
                    # If the sentence is empty, skip it
                    if not sentence.strip():
                        continue

                    # Check if the sentence needs correction
                    for incorrect, correct in corrections.items():
                        if incorrect in sentence:
                            start, end = diff_strings(incorrect, correct)
                            incorrect_part = incorrect[start:end]
                            corrected_part = correct[start:end]

                            # Split the sentence around the incorrect part
                            before, _, after = sentence.partition(incorrect_part)

                            # Add the part before the incorrect part
                            new_paragraph.add_run(before)

                            # Add the incorrect part with strike-through
                            run = new_paragraph.add_run(incorrect_part)
                            run.font.strike = True

                            # Add the corrected part in brackets and color it red
                            corrected_run = new_paragraph.add_run(" (" + corrected_part + ")")
                            corrected_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # RGB for red

                            # Add the part after the incorrect part
                            new_paragraph.add_run(after + ".")

                            break  # Break the loop once a correction is applied

                    else:
                        # If the sentence doesn't need correction, add it without any modification
                        new_paragraph.add_run(sentence + ".")

                
                # Add improvement suggestions after each paragraph
                for original, improvement in improvements.items():
                    if original in paragraph.text:

                        # Check if "建议改为" is in the improvement
                        if "建议改为" not in improvement:
                            # If not, prepend "建议改为"
                            improvement = "建议改为" + improvement

                        # Split on "建议改为" to get the suggestion
                        suggestion = improvement.split("建议改为")[1]

                        # Create a new paragraph for the suggestion
                        improved_paragraph = corrected_doc.add_paragraph()
                        
                        # Add the original sentence with underline
                        run = improved_paragraph.add_run(original + ". ")
                        run.font.underline = True

                        # Add "建议改为"
                        improved_paragraph.add_run("建议改为")

                        # Add the suggestion in blue
                        run = improved_paragraph.add_run(suggestion)
                        run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)  # RGB for blue



            # Add IELTS comments

            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)
            # title_line = corrected_doc.add_paragraph("提分指南:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('提分指南', 1) 
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(development)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
            # Add IELTS rating

            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)            
            # title_line = corrected_doc.add_paragraph("习作评分:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('习作评分', 1) 
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(rating)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Add a model essay

            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)   
            # title_line = corrected_doc.add_paragraph("满分范文:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('满分范文', 1) 
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(sample)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Add an analysis of the model essay

            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)   
            # title_line = corrected_doc.add_paragraph("范文点评:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('范文点评', 1) 
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(sample_rating)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Add details for this essay question

            run = corrected_doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)   
            # title_line = corrected_doc.add_paragraph("思路建议:")
            # run = title_line.runs[0]
            # run.font.size = Pt(20)
            # run.font.color.rgb = RGBColor(0x64, 0xAA, 0xE7)
            # title_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            corrected_doc.add_heading('思路建议', 1) 
            paragraph = corrected_doc.add_paragraph()
            run = paragraph.add_run(ideas)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
                
            # Save the modified document
            filename = '/home/ec2-user/mysite/static/dlFile/ToeflTask2_modified.docx'
            corrected_doc.save(filename)

            # 返回相对路径
            return filename
        # Call your function that generates the file
        full_filepath = modify_word_document(essay)
        update_generate_count(-40)

        directory, filename = os.path.split(full_filepath)
        print(f"Trying to send file: {filename} from directory: {directory}")

        # Create a response with the file
        response = make_response(send_file(full_filepath, as_attachment=True))
        response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        return response

    except Exception as e:
        print(f"Error: {str(e)}")
        traceback.print_exc()
        return jsonify(status=500, message="服务器错误。", data={"error": str(e)})

#------------------ generate_PDF --------------------------#
